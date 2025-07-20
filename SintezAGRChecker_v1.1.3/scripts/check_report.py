import datetime
import subprocess
import site
import sys
import shutil
import os

import bpy

from . import utills
from . import ui_utills
from . import os_utils
from . import logger


def get_dont_checked_nums(categories):
        nums = []
        for cat in categories:
            for item in cat.collection:
                if item.check_state == utills.CHECK_STATE_ITEMS[0]:
                    nums.append(item.req_num)
        return nums

def _generate_report(context):
    props = context.scene.agr_scene_properties
    categories_lp_hp = [props.checklist_lp_props.categories, props.checklist_hp_props.categories]
    msg_for_text_editor = ""
    msg_for_text_editor += "--------Отчет об ошибках--------\n\n"
    msg_for_text_editor += f"Address: {props.project_data_address}\n"
    dt = datetime.datetime.now()
    day = str(dt.day).rjust(2, "0")
    month = str(dt.month).rjust(2, "0")
    hour = str(dt.hour).rjust(2, "0")
    minute = str(dt.minute).rjust(2, "0")
    msg_for_text_editor += f"Дата проверки: {day}.{month}.{dt.year} {hour}:{minute}\n"
    if props.check_author:
        msg_for_text_editor += f"Проверяющий: {props.check_author}\n\n"
    else:
        msg_for_text_editor += "\n"

    if props.has_lowpoly:
        progress = round(props.checklist_lp_props.progress_all * 100)
        msg_for_text_editor += f"НПМ модель соответствует требованиям на {progress}%\n"
        nums = get_dont_checked_nums(props.checklist_lp_props.categories)
        if nums:
            msg_for_text_editor += f"НПМ. Проверены все пункты, кроме - {', '.join(nums)}\n"
        msg_for_text_editor += "\n"
    if props.has_highpoly:
        progress = round(props.checklist_hp_props.progress_all * 100)
        msg_for_text_editor += f"ВПМ модель соответствует требованиям на {progress}%\n"
        nums = get_dont_checked_nums(props.checklist_hp_props.categories)
        if nums:
            msg_for_text_editor += f"ВПМ. Проверены все пункты, кроме - {', '.join(nums)}\n"
        msg_for_text_editor += "\n"
    
    msg_for_text_editor += f"Нумерация пунктов взята из требований, сокращенная запись - Таблица.Пункт.Подпункт\n"
    msg_for_text_editor += f"Ссылка на требования: " + r"https://www.mos.ru/mka/function/gosudarstvennye-uslugi/svidetelstvo-ob-utverzhdenii-agr/" + "\n"
    msg_for_text_editor += f"Отчет подготовлен с помощью плагина Sintez AGR Checker: " + r"https://sintez.space/" + "\n"
    msg_for_text_editor += "\n\n"

    msg_for_text_editor_2 = msg_for_text_editor

    for i, categories in enumerate(categories_lp_hp):
        if i == 0:
            msg_for_text_editor += "--------Низкополигональная модель--------\n\n"
            msg_for_text_editor_2 += "--------Низкополигональная модель--------\n\n"
        else:
            msg_for_text_editor += "--------Высокополигональная модель--------\n\n"
            msg_for_text_editor_2 += "--------Высокополигональная модель--------\n\n"
        errr_num = 1
        for cat in categories:
            for item in cat.collection:
                if item.check_state == utills.CHECK_STATE_ITEMS[2]:
                    msg_for_text_editor += f"----Ошибка {errr_num}. Пункт {item.req_num}. {item.name}\n"
                    msg_for_text_editor_2 += f"Пункт {item.req_num}\n"
                    errr_num += 1
                    if not item.user_description:
                        item.user_description = item.description
                        item.user_description.replace("\n", "\n    ")
                    msg_for_text_editor += f"----Описание пункта:\n{item.user_description}\n"
                    msg_for_text_editor_2 += f"{item.user_description}\n"
                    if item.user_comment:
                        msg_for_text_editor_2 += f"{item.user_comment}\n"
                    if item.auto:
                        msg_for_text_editor += f"----Автоматичсеские ошибки:\n"
                        msg_for_text_editor += item.errors_text
                        msg_for_text_editor_2 += item.errors_text
                    msg_for_text_editor += f"----Комментарий к ошибке:\n{item.user_comment}\n"
                    msg_for_text_editor += "\n"
                    msg_for_text_editor_2 += "\n"
                    if item.user_comment:
                        msg_for_text_editor += "\n"
    return msg_for_text_editor, msg_for_text_editor_2

def ensure_import_docx():
    try:
        user_site_pkgs = site.getusersitepackages()
        if user_site_pkgs not in sys.path:
            sys.path.append(user_site_pkgs)
        import docx
        return True
    except Exception as e:
        try:
            logger.add('installing package docx')
            py_exec = sys.executable
            subprocess.call([str(py_exec), "-m", "ensurepip", "--user"])
            subprocess.call([str(py_exec), "-m", "pip", "install", "--upgrade", "pip"])
            subprocess.run([str(py_exec), "-m", "pip", "install", "python-docx"])
            import docx
            return True
        
        except Exception as e:
            logger.add_error(e, f"Не удалось установить библиотеку docx (для создания отчета в MS Word)")
            return False

def save_report_txt_file(context, bl_operator):
    if not os_utils.check_models_path():
        return
    
    ui_utills.update_checks_from_text_editor(context)
    report_msg, report_msg_2 = _generate_report(context)

    report_path = os_utils.get_checks_data_dir()
    
    dt = datetime.datetime.now()
    day = str(dt.day).rjust(2, "0")
    month = str(dt.month).rjust(2, "0")
    props = bpy.context.scene.agr_scene_properties

    lp_suf = "_НПМ" if props.has_lowpoly else ""
    hp_suf = "_ВПМ" if props.has_highpoly else ""

    with open(os.path.join(report_path, f"{props.project_data_address}_CheckReport_data.txt"), "w") as file:
        file.write(report_msg)
    with open(os.path.join(report_path, f"{dt.year}{month}{day}_{props.project_data_address}_Отчет{lp_suf}{hp_suf}.txt"), "w") as file:
        file.write(report_msg_2)
    save_docx_report(context, bl_operator)

def save_docx_report(context, bl_operator):
    if not os_utils.check_models_path():
        return
    if not ensure_import_docx():
        return
    
    from docx import Document
    from docx.shared import Cm, Pt

    dt = datetime.datetime.now()
    day = str(dt.day).rjust(2, "0")
    month = str(dt.month).rjust(2, "0")
    props = bpy.context.scene.agr_scene_properties
    categories_lp_hp = [props.checklist_lp_props.categories, props.checklist_hp_props.categories]

    src = os.path.join(os_utils.get_documents_dir(), "report_template.docx")
    lp_suf = "_НПМ" if props.has_lowpoly else ""
    hp_suf = "_ВПМ" if props.has_highpoly else ""
    report_file_name = f"{dt.year}{month}{day}_{props.project_data_address}_Отчет{lp_suf}{hp_suf}.docx"
    dst = os.path.join(os_utils.get_checks_data_dir(), report_file_name)
    try:
        # doc.save(docx_path)
        shutil.copyfile(src, dst)
    except:
        bl_operator.report({'WARNING'}, f"Закройте файл для сохрвнения отчета!")
    
    doc = Document(dst)

    style = doc.styles['Normal']
    # style.font.name = 'Arial'
    # style.font.size = Pt(14)
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    # p = doc.add_paragraph().add_run("Отчет об ошибках\n")
    # p.font.bold = True
    doc.add_paragraph().add_run(f"Address: {props.project_data_address}")
    # doc.add_paragraph().add_run(f"Дата проверки: {day}.{month}.{dt.year} {hour}:{minute}")
    doc.add_paragraph().add_run(f"Дата проверки: {day}.{month}.{dt.year}")
    # zip_path = ""
    # for root, dirs, files in os.walk(os_utils._get_project_path()):
    #     for file in files:
    #         if file.endswith(".zip") and file.startswith("SM_") or file[:4].isdigit():
    #             zip_path = os.path.join(root, file)
    #             break
    # if zip_path:
    #     zip_dt = datetime.datetime.fromtimestamp(os.path.getmtime(zip_path))
    #     zip_day = str(zip_dt.day).rjust(2, "0")
    #     zip_month = str(zip_dt.month).rjust(2, "0")
    #     doc.add_paragraph().add_run(f"Дата изменения архивов: {zip_day}.{zip_month}.{zip_dt.year}")
    # os.path.getmtime(path)
    if props.check_author:
        doc.add_paragraph().add_run(f"Проверяющий: {props.check_author}\n")
    else:
        doc.add_paragraph()
        pass

    if props.has_lowpoly:
        progress = round(props.checklist_lp_props.progress_all * 100)
        doc.add_paragraph().add_run(f"НПМ модель соответствует требованиям на {progress}%")
        nums = get_dont_checked_nums(props.checklist_lp_props.categories)
        if nums:
            doc.add_paragraph().add_run(f"НПМ. Проверены все пункты, кроме - {', '.join(nums)}")
        doc.add_paragraph()
    if props.has_highpoly:
        progress = round(props.checklist_hp_props.progress_all * 100)
        doc.add_paragraph().add_run(f"ВПМ модель соответствует требованиям на {progress}%")
        nums = get_dont_checked_nums(props.checklist_hp_props.categories)
        if nums:
            doc.add_paragraph().add_run(f"ВПМ. Проверены все пункты, кроме - {', '.join(nums)}")
        doc.add_paragraph()
    
    # doc.add_paragraph().add_run(f"Нумерация пунктов взята из требований, сокращенная запись - Таблица.Пункт.Подпункт")
    # doc.add_paragraph().add_run(f"Ссылка на требования: " + r"https://www.mos.ru/mka/function/gosudarstvennye-uslugi/svidetelstvo-ob-utverzhdenii-agr/" + "")
    # doc.add_paragraph().add_run(f"Отчет подготовлен с помощью плагина Sintez AGR Checker: " + r"https://sintez.space/" + "")
    # doc.add_paragraph()
    doc.add_paragraph()

    for i, categories in enumerate(categories_lp_hp):
        if i == 0:
            p = doc.add_paragraph().add_run("Низкополигональная модель. Ошибки")
            p.font.bold = True
            p = doc.add_paragraph().add_run(f"Нумерация пунктов взята из требований, сокращенная запись - Таблица.Пункт.Подпункт\n")
            p.font.italic = True
        else:
            p = doc.add_paragraph().add_run("Высокополигональная модель. Ошибки")
            p.font.bold = True
            p = doc.add_paragraph().add_run(f"Нумерация пунктов взята из требований, сокращенная запись - Таблица.Пункт.Подпункт\n")
            p.font.italic = True
        for cat in categories:
            for item in cat.collection:
                if item.check_state == utills.CHECK_STATE_ITEMS[2]:
                    p = doc.add_paragraph().add_run(f"Пункт {item.req_num}")
                    p.font.bold = True
                    if not item.user_description:
                        item.user_description = item.description
                        item.user_description.replace("\n", "\n    ")
                    p = doc.add_paragraph().add_run(f"{item.user_description}")
                    p.italic = True
                    if item.user_comment:
                        p = doc.add_paragraph().add_run(f"{item.user_comment}")
                    if item.auto:
                        p = doc.add_paragraph().add_run(item.errors_text)
                    for im_prop in item.error_images:
                        doc.add_picture(im_prop.full_path, width=Cm(16))
                    p = doc.add_paragraph()

    # document.add_paragraph().add_run(report_msg_2)
    # for cat in context.scene.agr_scene_properties.checklist_lp_props.categories:
    #     for item in cat.collection:
    #         for im_prop in item.error_images:
    #             document.add_picture(im_prop.full_path, width=Cm(15))

    # models_path = os_utils._get_project_path()
    # docx_path = models_path + f"{context.scene.agr_scene_properties.project_data_address}_CheckReport_MSword2.docx"
    try:
        # doc.save(docx_path)
        doc.save(dst)
    except:
        bl_operator.report({'WARNING'}, f"Закройте файл для сохрвнения отчета!")
    # pdf_path = models_path + f"{context.scene.agr_scene_properties.project_data_address}_CheckReport.pdf"
    # try_create_pdf(docx_path, pdf_path)

def try_create_pdf(path_docx, path_pdf):
    try:
        import os
        import comtypes.client
        import time

        wdFormatPDF = 17

        # absolute path is needed
        # be careful about the slash '\', use '\\' or '/' or raw string r"..."
        in_file = path_docx
        out_file = path_pdf

        # create COM object
        word = comtypes.client.CreateObject('Word.Application')
        # key point 1: make word visible before open a new document
        word.Visible = True
        # key point 2: wait for the COM Server to prepare well.
        time.sleep(3)

        # convert docx file 1 to pdf file 1
        # doc = word.Documents.Open(in_file) # open docx file 1
        # doc.SaveAs(out_file, FileFormat=wdFormatPDF) # conversion
        # doc.Close() # close docx file 1
        word.Visible = False
        # convert docx file 2 to pdf file 2
        doc = word.Documents.Open(in_file) # open docx file 2
        doc.SaveAs(out_file, FileFormat=wdFormatPDF) # conversion
        doc.Close() # close docx file 2   
        word.Quit() # close Word Application 
    except Exception as e:
        logger.add_error(e)

def load_report(context):
    path = bpy.path.abspath(context.scene.agr_scene_properties.load_report_path)
    models_path = bpy.path.abspath(context.scene.agr_scene_properties.path)
    logger.add(f"models path = {models_path}")
    logger.add(f"load report path = {path}")
    if not path.endswith(".txt"):
        return
    failed_lp_req_nums = []
    failed_hp_req_nums = []
    failed_nums = failed_lp_req_nums
    not_checked_lp_req_nums = []
    not_checked_hp_req_nums = []
    msg_for_text_editor = ""
    with open(path) as file:
        for line in file.readlines():
            msg_for_text_editor += line
            if "НПМ. Проверены все пункты, кроме" in line:
                not_checked_lp_req_nums = line.split(" - ")[1].strip().split(", ")
            elif "ВПМ. Проверены все пункты, кроме" in line:
                not_checked_hp_req_nums = line.split(" - ")[1].strip().split(", ")
            elif "Высокополигональная модель" in line:
                failed_nums = failed_hp_req_nums
            elif line.startswith("----Ошибка"):
                failed_nums.append(line.split(" ")[3].strip("."))
    
    categories_lp_hp = [(context.scene.agr_scene_properties.checklist_lp_props.categories, failed_lp_req_nums, not_checked_lp_req_nums),
                        (context.scene.agr_scene_properties.checklist_hp_props.categories, failed_hp_req_nums, not_checked_hp_req_nums)]
    for pair in categories_lp_hp:
        categories = pair[0]
        failed_nums = pair[1]
        not_checked = pair[2]
        for cat in categories:
            for item in cat.collection:
                if item.req_num in failed_nums:
                    item.check_state = utills.CHECK_STATE_ITEMS[2]
                    item.check = False
                elif item.req_num not in not_checked:
                    item.check_state = utills.CHECK_STATE_ITEMS[1]
                    item.check = True
    
    text_file = ui_utills.get_text_editor_text(context)

    text_file.write(msg_for_text_editor)
    ui_utills.update_checks_from_text_editor(context)